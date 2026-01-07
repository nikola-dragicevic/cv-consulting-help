# scripts/parse_cv_pdf.py
import fitz  # PyMuPDF
import docx
from typing import Optional, Tuple


def extract_text_from_pdf(file_path: str) -> Tuple[str, bool]:
    """
    Extract raw text from a PDF CV and detect if it contains images.
    Returns: (text, has_picture_bool)

    Notes:
    - page.get_text() ignores images automatically.
    - page.get_images(full=True) detects embedded images on the page.
    """
    text_parts = []
    has_picture = False

    try:
        with fitz.open(file_path) as doc:
            for page in doc:
                # Extract text
                page_text = page.get_text() or ""
                if page_text:
                    text_parts.append(page_text)

                # Detect images (only until we find one)
                if not has_picture:
                    try:
                        images = page.get_images(full=True)
                        if images:
                            has_picture = True
                    except Exception:
                        # Don't fail parsing because image detection failed
                        pass

    except Exception as e:
        print(f"❌ Error reading PDF: {e}")

    return "".join(text_parts).strip(), has_picture


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file.
    Includes:
    - paragraphs
    - tables (common in Swedish CV templates)
    - headers/footers (some CVs put name/contact there)
    """
    parts = []

    try:
        doc = docx.Document(file_path)

        # Paragraphs
        for para in doc.paragraphs:
            t = (para.text or "").strip()
            if t:
                parts.append(t)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    ct = (cell.text or "").strip()
                    if ct:
                        # avoid repeating same cell text due to merged cells
                        cells.append(ct)
                if cells:
                    parts.append(" | ".join(cells))

        # Headers/Footers
        for section in doc.sections:
            header = section.header
            footer = section.footer

            if header:
                for p in header.paragraphs:
                    t = (p.text or "").strip()
                    if t:
                        parts.append(t)

            if footer:
                for p in footer.paragraphs:
                    t = (p.text or "").strip()
                    if t:
                        parts.append(t)

    except Exception as e:
        print(f"❌ Error reading DOCX: {e}")

    return "\n".join(parts).strip()


def summarize_cv_text(cv_text: str) -> Optional[str]:
    """
    Cleans CV text and removes database-breaking characters.
    Keeps Unicode (Swedish chars).
    Preserves structure via newlines while removing empty lines.
    """
    if not cv_text:
        return None

    # ✅ CRITICAL: Remove Null bytes that crash Postgres
    cv_text = cv_text.replace("\x00", "")

    # Normalize line endings
    cv_text = cv_text.replace("\r", "\n")

    # Strip whitespace per line, remove empty lines
    lines = [line.strip() for line in cv_text.splitlines()]
    cleaned_lines = [line for line in lines if line]

    cleaned = "\n".join(cleaned_lines).strip()
    return cleaned if cleaned else None
